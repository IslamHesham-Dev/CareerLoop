from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BACKEND_ROOT = Path(__file__).resolve().parent.parent
PROJECT_ROOT = BACKEND_ROOT.parent
CATALOG_PATH = BACKEND_ROOT / "content" / "cms_catalog.json"
MARKDOWN_PATH = BACKEND_ROOT / "content" / "transcript_intake_template.md"
DOCX_PATH = PROJECT_ROOT / "docs" / "CareerLoop_CMS_Transcript_Intake.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(92, 101, 115)
LIGHT_FILL = "F4F6F9"
BORDER = "D9E1EA"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeatable_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def set_cell_margins(cell, *, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(6.5)
    table_pr = table._tbl.tblPr

    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")

    table_indent = table_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    grid.append(grid_col)

    for cell in table.row_cells(0):
        tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            cell._tc.get_or_add_tcPr().append(tc_w)
        tc_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
        tc_w.set(qn("w:type"), "dxa")
        set_cell_margins(cell)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_section(section) -> None:
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.text = ""
    left = header_paragraph.add_run("CareerLoop CMS")
    set_run_font(left, size=9, color=MUTED, bold=True)
    right = header_paragraph.add_run("  |  Transcript intake")
    set_run_font(right, size=9, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = ""
    add_page_field(footer_paragraph)


def add_placeholder(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, LIGHT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("[PASTE TRANSCRIPT OR DETAILED SUMMARY HERE]")
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    table_borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    if table_borders is not None:
        for edge in table_borders:
            edge.set(qn("w:color"), BORDER)
            edge.set(qn("w:sz"), "6")
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build_markdown(catalog: dict) -> str:
    lines = [
        "# CareerLoop CMS Transcript Intake",
        "",
        "Paste a transcript or detailed summary between the matching START and "
        "END markers. Do not edit the Drive file ID.",
        "",
    ]
    for course in catalog["courses"]:
        lines.extend(
            [
                f"## {course['title']} ({course['catalog_code']})",
                "",
            ]
        )
        for item in course["items"]:
            lines.extend(
                [
                    f"### {item['title']}",
                    "",
                    f"- Drive file ID: `{item['id']}`",
                    f"- Type: {item['content_type']}",
                    f"- Video: {item['drive_url']}",
                    "",
                    f"<!-- TRANSCRIPT START: {item['id']} -->",
                    "[PASTE TRANSCRIPT OR DETAILED SUMMARY HERE]",
                    f"<!-- TRANSCRIPT END: {item['id']} -->",
                    "",
                ]
            )
    return "\n".join(lines).rstrip() + "\n"


def build_docx(catalog: dict) -> None:
    document = Document()
    section = document.sections[0]
    configure_section(section)
    set_repeatable_styles(document)

    cover_space = document.add_paragraph()
    cover_space.paragraph_format.space_after = Pt(64)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("CAREERLOOP CMS")
    set_run_font(run, size=11, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Transcript Intake")
    set_run_font(run, size=30, color=INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run(
        "A source-grounded template for lecture and tutorial transcripts"
    )
    set_run_font(run, size=14, color=DARK_BLUE)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_after = Pt(70)
    total = sum(len(course["items"]) for course in catalog["courses"])
    run = metadata.add_run(
        f"{len(catalog['courses'])} courses  |  {total} videos  |  "
        "Google Drive inventory"
    )
    set_run_font(run, size=10.5, color=MUTED, bold=True)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(4)
    run = note.add_run(
        "Keep each Drive file ID unchanged. Paste text only inside its "
        "matching transcript field."
    )
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    document.add_page_break()
    document.add_heading("How to fill this document", level=1)
    instructions = [
        (
            "1. Open the linked video",
            "Use the Drive link under the video title to confirm the recording.",
        ),
        (
            "2. Paste the source text",
            "Replace the bracketed placeholder with the transcript or a "
            "detailed summary. Keep the video heading and Drive file ID.",
        ),
        (
            "3. Preserve uncertainty",
            "If audio is unclear, mark the passage as [inaudible] instead of "
            "guessing. Keep code, formulas, and terminology as close to the "
            "recording as possible.",
        ),
        (
            "4. Return the completed file",
            "CareerLoop will split entries by Drive file ID, create individual "
            "transcript sources, and make them available to the study agent.",
        ),
    ]
    for heading, body in instructions:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        first = paragraph.add_run(f"{heading}. ")
        set_run_font(first, bold=True, color=INK)
        second = paragraph.add_run(body)
        set_run_font(second)

    for course_index, course in enumerate(catalog["courses"]):
        if course_index == 0:
            document.add_page_break()
        else:
            document.add_section(WD_SECTION.NEW_PAGE)
            configure_section(document.sections[-1])

        document.add_heading(
            f"{course['title']} ({course['catalog_code']})",
            level=1,
        )
        summary = document.add_paragraph()
        summary.paragraph_format.space_after = Pt(10)
        run = summary.add_run(
            f"{len(course['items'])} videos. {course['description']}"
        )
        set_run_font(run, color=MUTED)

        for item in course["items"]:
            heading = document.add_heading(item["title"], level=3)
            heading.paragraph_format.keep_with_next = True

            metadata_paragraph = document.add_paragraph()
            metadata_paragraph.paragraph_format.space_after = Pt(4)
            metadata_paragraph.paragraph_format.keep_with_next = True
            label = metadata_paragraph.add_run(
                f"{item['content_type'].title()}  |  Drive file ID: "
                f"{item['id']}  |  "
            )
            set_run_font(label, size=9.5, color=MUTED)
            add_hyperlink(
                metadata_paragraph,
                "Open Drive video",
                item["drive_url"],
            )

            transcript_label = document.add_paragraph()
            transcript_label.paragraph_format.space_after = Pt(3)
            transcript_label.paragraph_format.keep_with_next = True
            run = transcript_label.add_run("Transcript / detailed summary")
            set_run_font(run, size=10, color=INK, bold=True)
            add_placeholder(document)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)


def validate_outputs(catalog: dict) -> None:
    expected = sum(len(course["items"]) for course in catalog["courses"])
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    if markdown.count("<!-- TRANSCRIPT START:") != expected:
        raise RuntimeError("Markdown START marker count is incorrect.")
    if markdown.count("<!-- TRANSCRIPT END:") != expected:
        raise RuntimeError("Markdown END marker count is incorrect.")

    document = Document(DOCX_PATH)
    placeholders = sum(
        paragraph.text == "[PASTE TRANSCRIPT OR DETAILED SUMMARY HERE]"
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    hyperlinks = len(document.element.body.xpath(".//w:hyperlink"))
    if len(document.tables) != expected:
        raise RuntimeError("DOCX placeholder table count is incorrect.")
    if placeholders != expected:
        raise RuntimeError("DOCX placeholder text count is incorrect.")
    if hyperlinks != expected:
        raise RuntimeError("DOCX Drive hyperlink count is incorrect.")

    for table in document.tables:
        table_width = table._tbl.tblPr.find(qn("w:tblW"))
        table_indent = table._tbl.tblPr.find(qn("w:tblInd"))
        cell_width = table.cell(0, 0)._tc.get_or_add_tcPr().find(
            qn("w:tcW")
        )
        if table_width is None or table_width.get(qn("w:w")) != "9360":
            raise RuntimeError("DOCX table width audit failed.")
        if table_indent is None or table_indent.get(qn("w:w")) != "120":
            raise RuntimeError("DOCX table indent audit failed.")
        if cell_width is None or cell_width.get(qn("w:w")) != "9360":
            raise RuntimeError("DOCX cell width audit failed.")


def main() -> None:
    catalog = json.loads(CATALOG_PATH.read_text(encoding="utf-8"))
    MARKDOWN_PATH.write_text(build_markdown(catalog), encoding="utf-8")
    build_docx(catalog)
    validate_outputs(catalog)
    print(MARKDOWN_PATH)
    print(DOCX_PATH)
    print("Validated 153 transcript placeholders and Drive hyperlinks.")


if __name__ == "__main__":
    main()
